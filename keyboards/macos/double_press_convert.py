#!/usr/bin/env python3
r"""
double_press_convert.py
Convert LSD/CRULP double-press Arabic text to proper Unicode.

When typing on an LSD keyboard, secondary characters are entered by pressing
the same key twice quickly (سس → ے, جج → چھے, etc.).  Old Word documents
may contain these literal repeated-character sequences instead of the
correct Unicode characters.  This tool replaces them in-place.

Modes
-----
  python3 double_press_convert.py input.docx [output.docx]   # Word doc
  python3 double_press_convert.py --crulp input.docx [out.docx]
  echo "سسضض" | python3 double_press_convert.py              # plain text
  cat file.txt | python3 double_press_convert.py --crulp

If output path is omitted the input file is overwritten (docx mode).

Substitution tables (LSD layout, default)
------------------------------------------
  جج → چھے    سس → ے    ضض → ٹ    ثث → پ    هه → ھ    حح → چ
  يي → ئ     اا → اٰ    كك → گ    طط → ں    رر → ڑ    ةة → ۃ
  دد → ڈ     ظظ → ہ

CRULP additions/overrides
--------------------------
  عع → غ     تت → ٹ    ہہ → ھ    زز → ذ    شش → ض    نن → ں
  حح → خ     (رر → ڑ and دد → ڈ same as LSD)
  Also applies primary-char swaps: ه→ہ  ك→ک  ة→ۃ  ي→ی
"""

import sys
import re
from typing import Sequence

try:
    import docx
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Substitution tables
# ---------------------------------------------------------------------------

# LSD double-press: each entry is (repeated_char, secondary).
# Process longest-output entries first so جج→چھے doesn't become چ+extra.
_LSD_DOUBLES = [
    ("ج", "چھے"),
    ("ض", "ٹ"),  ("ث", "پ"),  ("ه", "ھ"),  ("ح", "چ"),  ("س", "ے"),
    ("ي", "ئ"),  ("ا", "اٰ"), ("ك", "گ"),  ("ط", "ں"),  ("ر", "ڑ"),
    ("ة", "ۃ"),  ("د", "ڈ"),  ("ظ", "ہ"),
]

_CRULP_DOUBLES = [
    ("ع", "غ"),  ("ر", "ڑ"),  ("ت", "ٹ"),  ("ح", "خ"),  ("د", "ڈ"),
    ("ہ", "ھ"),  ("ه", "ھ"),  ("ز", "ذ"),  ("ش", "ض"),  ("ن", "ں"),
    ("ج", "چھے"),
    ("ض", "ٹ"),  ("ث", "پ"),  ("س", "ے"),  ("ي", "ئ"),  ("ا", "اٰ"),
    ("ك", "گ"),  ("ط", "ں"),  ("ر", "ڑ"),  ("ة", "ۃ"),  ("ظ", "ہ"),
]

# CRULP primary-char swaps (applied after double-press substitution)
_CRULP_PRIMARY = {
    "ه": "ہ",   # he → he goal
    "ك": "ک",   # kaf → high kaf
    "ة": "ۃ",   # te marbuta → te marbuta with ring
    "ي": "ی",   # arabic yeh → farsi yeh
}


def _build_pattern(doubles):
    """Compile a single regex that matches all double-press pairs."""
    alts = [re.escape(c + c) for c, _ in doubles]
    return re.compile("|".join(alts))


def _make_replacer(doubles):
    lookup = {c + c: s for c, s in doubles}
    pattern = _build_pattern(doubles)

    def replace(text: str) -> str:
        return pattern.sub(lambda m: lookup[m.group()], text)

    return replace


_lsd_replace = _make_replacer(_LSD_DOUBLES)
_crulp_replace = _make_replacer(_CRULP_DOUBLES)


def convert_text(text: str, crulp: bool = False) -> str:
    """Apply double-press substitutions to a plain string."""
    if crulp:
        text = _crulp_replace(text)
        for src, dst in _CRULP_PRIMARY.items():
            text = text.replace(src, dst)
    else:
        text = _lsd_replace(text)
    return text


# ---------------------------------------------------------------------------
# Word document processing
# ---------------------------------------------------------------------------

def _all_paragraphs(doc):
    """Yield every paragraph in the document (body + tables + headers/footers)."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf is not None:
                yield from hf.paragraphs


def _fix_paragraph(para, crulp: bool):
    """
    Apply substitutions within each run, then handle cross-run boundaries.

    Cross-run pass: if run[i] ends with X and run[i+1] starts with X and
    X+X is a double-press pair, merge those two characters and substitute.
    """
    runs = para.runs
    if not runs:
        return

    # Within-run substitution
    for run in runs:
        if run.text:
            run.text = convert_text(run.text, crulp)

    # Cross-run boundary pass (one pass is sufficient since substitution
    # consumes both copies of the character)
    doubles = _CRULP_DOUBLES if crulp else _LSD_DOUBLES
    lookup = {c + c: s for c, s in doubles}

    for i in range(len(runs) - 1):
        a, b = runs[i], runs[i + 1]
        if not a.text or not b.text:
            continue
        pair = a.text[-1] + b.text[0]
        if pair in lookup:
            replacement = lookup[pair]
            a.text = a.text[:-1] + replacement
            b.text = b.text[1:]

    # CRULP primary swaps after cross-run pass
    if crulp:
        for run in runs:
            if run.text:
                for src, dst in _CRULP_PRIMARY.items():
                    run.text = run.text.replace(src, dst)


def convert_docx(src_path: str, dst_path: str, crulp: bool = False):
    doc = docx.Document(src_path)
    changed = 0
    for para in _all_paragraphs(doc):
        before = "".join(r.text for r in para.runs)
        _fix_paragraph(para, crulp)
        after = "".join(r.text for r in para.runs)
        if before != after:
            changed += 1
    doc.save(dst_path)
    return changed


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv: Sequence[str]):
    args = list(argv[1:])
    crulp = "--crulp" in args
    if crulp:
        args.remove("--crulp")

    if not args and not sys.stdin.isatty():
        # Plain-text filter mode
        text = sys.stdin.read()
        sys.stdout.write(convert_text(text, crulp))
        return

    if not args:
        print(__doc__)
        sys.exit(0)

    src = args[0]
    dst = args[1] if len(args) > 1 else src

    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed.  Run: pip install python-docx",
              file=sys.stderr)
        sys.exit(1)

    if not src.lower().endswith(".docx"):
        print(f"Error: expected a .docx file, got: {src}", file=sys.stderr)
        sys.exit(1)

    n = convert_docx(src, dst, crulp)
    layout = "CRULP" if crulp else "LSD"
    action = "overwritten" if dst == src else f"saved to {dst!r}"
    print(f"{layout}: {n} paragraph(s) changed — {action}")


if __name__ == "__main__":
    main(sys.argv)

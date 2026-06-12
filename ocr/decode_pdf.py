#!/usr/bin/env python3
"""Decode a legacy-font LSD PDF (AL-FATEMI / AL-KANZ ...) to clean Unicode text.

The end-to-end "give me the document back" pipeline, no manual steps:

  glyph image-matching (legacy_decode, per-page cache)
    -> paragraph merge (RTL margin heuristic)
    -> post-pass (postpass.postprocess: template-confusion fixes, gated
       double-press + emphasis-doubling collapse via lsd-normalize, digit
       un-mirroring, heh-goal recovery)
    -> .txt (one paragraph per line, blank line between pages)
    -> optional .docx (--docx; needs python-docx; RTL paragraphs in
       FatemiMaqala, page breaks aligned with the PDF)

Typical run (docker, worktree at /work, PDFs in /data):

  docker run --rm -v ~/Documents/font-fatemi-ocr:/work -v /tmp/job:/data \
      -v ~/Documents/lsd-normalize:/opt/lsd-normalize lsd-ocr:latest \
      python3 /work/ocr/decode_pdf.py /data/book.pdf --verify-page 11

Always eyeball the --verify-page sheet (PDF crop above the same text re-rendered
in the reference font) before trusting a new document — new font subsets can
carry new template confusions; add them to postpass.WORD_FIXES.
"""
from __future__ import annotations

import argparse
import os
import pathlib
import sys
import time

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))

import fitz  # noqa: E402
from legacy_decode import LegacyDecoder, decode_page_lines, extract_page_faces  # noqa: E402
import postpass  # noqa: E402

MARGIN_TOL = 18.0  # pt; how far from the margin still counts as "reaching" it


def paragraphs(lines: list[tuple[str, tuple]]) -> list[str]:
    """Merge (text, bbox) visual lines into paragraphs (RTL-aware): a new
    paragraph starts when a line is indented at the right margin, or the
    previous line stopped short of the left margin."""
    if not lines:
        return []
    right = max(bb[2] for _, bb in lines)
    left = min(bb[0] for _, bb in lines)
    paras, cur, prev_bb = [], [], None
    for text, bb in lines:
        starts_new = (
            prev_bb is None
            or bb[2] < right - MARGIN_TOL
            or prev_bb[0] > left + MARGIN_TOL
        )
        if starts_new and cur:
            paras.append(" ".join(cur))
            cur = []
        cur.append(text)
        prev_bb = bb
    if cur:
        paras.append(" ".join(cur))
    return paras


def make_verify(page, decoded, ref_path: str, out_png: str, dpi: int = 200) -> None:
    """Stack each decoded line's PDF crop above the same text re-rendered in
    the reference font; if decoding is right the rows read the same."""
    from PIL import Image, ImageDraw
    from render import render_line

    scale = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), colorspace=fitz.csGRAY)
    img = Image.frombytes("L", (pix.width, pix.height), pix.samples)
    rows = []
    for text, bb in decoded:
        box = (max(0, int(bb[0] * scale) - 4), max(0, int(bb[1] * scale) - 4),
               min(img.width, int(bb[2] * scale) + 4),
               min(img.height, int(bb[3] * scale) + 4))
        if box[2] <= box[0] or box[3] <= box[1]:
            continue
        crop = img.crop(box)
        try:
            rendered = render_line(text, ref_path, 40)
        except Exception:
            rendered = Image.new("L", (crop.width, 48), 255)
        rows.append((crop, rendered))
    if not rows:
        return
    pad, gap = 10, 6
    W = max(max(c.width, r.width) for c, r in rows) + 2 * pad
    H = sum(c.height + r.height + gap + 18 for c, r in rows) + pad
    sheet = Image.new("L", (W, H), 255)
    draw = ImageDraw.Draw(sheet)
    y = pad
    for crop, rendered in rows:
        sheet.paste(crop, (W - crop.width - pad, y)); y += crop.height + gap
        sheet.paste(rendered, (W - rendered.width - pad, y)); y += rendered.height
        draw.line([(pad, y + 8), (W - pad, y + 8)], fill=200); y += 18
    sheet.save(out_png)
    print(f"  verify sheet -> {out_png} ({len(rows)} lines)", flush=True)


def write_docx(pages: list[str], out_path: str, font: str = "FatemiMaqala") -> None:
    """RTL docx, one PDF page per docx page."""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError:
        raise SystemExit("--docx needs python-docx (pip install python-docx)")

    d = docx.Document()
    style = d.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for attr in ("w:cs", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), font)

    for i, page in enumerate(pages):
        for para in page.split("\n"):
            if not para.strip():
                continue
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ppr = p._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi"); bidi.set(qn("w:val"), "1"); ppr.append(bidi)
            run = p.add_run(para.strip())
            rpr2 = run._r.get_or_add_rPr()
            rtl = OxmlElement("w:rtl"); rtl.set(qn("w:val"), "1"); rpr2.append(rtl)
            cs = OxmlElement("w:szCs"); cs.set(qn("w:val"), "28"); rpr2.append(cs)
        if i < len(pages) - 1:
            br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
            d.add_paragraph().add_run()._r.append(br)
    d.save(out_path)
    print(f"Wrote {out_path}", flush=True)


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("pdf", help="legacy-font PDF to decode")
    ap.add_argument("--ref", action="append", default=None,
                    help="reference font path(s); default: FatemiMaqala + Kanz "
                         "al-Marjaan (repeat for several; a single same-family "
                         "face decodes a touch cleaner than the combined set)")
    ap.add_argument("--out", default=None, help="output .txt (default: <pdf>.txt)")
    ap.add_argument("--docx", action="store_true", help="also write <out>.docx")
    ap.add_argument("--raw", action="store_true",
                    help="skip the post-pass (debugging the decoder itself)")
    ap.add_argument("--pages", default=None,
                    help="1-based page range 'a-b' or single page (default all)")
    ap.add_argument("--verify-page", type=int, default=0,
                    help="1-based page to dump a verify sheet for")
    args = ap.parse_args()

    out = args.out or str(pathlib.Path(args.pdf).with_suffix(".txt"))

    t0 = time.time()
    print(f"Building reference templates ...", flush=True)
    decoder = LegacyDecoder(reference_fonts=args.ref)
    ref_for_verify = decoder.ref_paths[-1]
    print(f"  {len(decoder._labels)} templates from {len(decoder.ref_paths)} "
          f"font(s) ({time.time()-t0:.1f}s)", flush=True)

    doc = fitz.open(args.pdf)
    if args.pages:
        a, _, b = args.pages.partition("-")
        prange = range(int(a) - 1, int(b or a))
    else:
        prange = range(doc.page_count)

    page_blocks, failed = [], []
    for pno in prange:
        page = doc[pno]
        faces, default_face, tmps = extract_page_faces(doc, page)
        try:
            # fresh per-page cache (created inside decode_page_lines)
            decoded = decode_page_lines(decoder, page, faces, default_face=default_face)
        except Exception as e:
            failed.append((pno + 1, repr(e)))
            decoded = []
        decoded = [(t.strip(), bb) for t, bb in decoded if t.strip()]
        block = "\n".join(paragraphs(decoded))
        if not args.raw:
            block = postpass.postprocess(block)
        page_blocks.append(block)

        if args.verify_page == pno + 1:
            make_verify(page, decoded, ref_for_verify,
                        str(pathlib.Path(out).with_suffix(".verify.png")))
        for t in tmps:
            os.unlink(t)
        if (pno + 1) % 10 == 0 or pno == prange[-1]:
            print(f"  page {pno+1}/{doc.page_count}  ({time.time()-t0:.0f}s elapsed)",
                  flush=True)

    with open(out, "w", encoding="utf-8") as fh:
        fh.write("\n\n".join(page_blocks) + "\n")
    print(f"Wrote {out}: {len(page_blocks)} page(s), "
          f"{sum(b.count(chr(10)) + 1 for b in page_blocks if b)} paragraph(s), "
          f"{time.time()-t0:.0f}s total", flush=True)
    if failed:
        print("FAILED pages:")
        for pno, err in failed:
            print(f"  page {pno}: {err}")
    if args.docx:
        write_docx(page_blocks, str(pathlib.Path(out).with_suffix(".docx")))


if __name__ == "__main__":
    main()
